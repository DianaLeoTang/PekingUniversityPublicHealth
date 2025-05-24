import re
import sys
import os
from docx import Document
import genanki
import random

def word_to_apkg(word_file_path, apkg_file_path, deck_name="学习卡片"):
    """
    将Word文件直接转换为Anki的.apkg文件
    
    参数:
    word_file_path: Word文件路径
    apkg_file_path: 输出APKG文件路径
    deck_name: 牌组名称
    """
    
    # 中文数字正则表达式模式
    pattern = r'^[一二三四五六七八九十]'
    
    try:
        # 读取Word文档
        print(f"正在读取Word文档: {word_file_path}")
        doc = Document(word_file_path)
        
        # 存储卡片数据
        cards_data = []
        current_title = ""
        current_content = []
        
        # 遍历文档中的每一段
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # 跳过空行
            if not text:
                continue
            
            # 检查是否是标题行（以中文数字开头）
            if re.match(pattern, text):
                # 如果之前有标题和内容，保存为一张卡片
                if current_title and current_content:
                    content_text = '<br>'.join(current_content).strip()
                    if content_text:  # 确保内容不为空
                        cards_data.append({
                            'front': current_title,
                            'back': content_text
                        })
                
                # 开始新的标题
                current_title = text
                current_content = []
            else:
                # 这是内容行，添加到当前内容中
                if current_title:  # 只有在有标题的情况下才添加内容
                    current_content.append(text)
        
        # 处理最后一张卡片
        if current_title and current_content:
            content_text = '<br>'.join(current_content).strip()
            if content_text:
                cards_data.append({
                    'front': current_title,
                    'back': content_text
                })
        
        if not cards_data:
            print("错误: 没有找到符合格式的卡片内容")
            return False
        
        print(f"找到 {len(cards_data)} 张卡片，开始生成APKG文件...")
        
        # 创建Anki模型（卡片模板）
        model = genanki.Model(
            random.randrange(1 << 30, 1 << 31),  # 随机模型ID
            '基础卡片模型',
            fields=[
                {'name': '正面'},
                {'name': '背面'},
            ],
            templates=[
                {
                    'name': '卡片1',
                    'qfmt': '''
                    <div class="card">
                        <div class="question">{{正面}}</div>
                    </div>
                    ''',
                    'afmt': '''
                    <div class="card">
                        <div class="question">{{正面}}</div>
                        <hr>
                        <div class="answer">{{背面}}</div>
                    </div>
                    ''',
                },
            ],
            css='''
            .card {
                font-family: "Microsoft YaHei", "Helvetica Neue", Helvetica, Arial, sans-serif;
                font-size: 18px;
                text-align: left;
                color: #333;
                background-color: #fff;
                padding: 20px;
                line-height: 1.6;
            }
            
            .question {
                font-weight: bold;
                font-size: 20px;
                color: #2c3e50;
                margin-bottom: 15px;
            }
            
            .answer {
                font-size: 16px;
                color: #34495e;
                margin-top: 15px;
            }
            
            hr {
                border: none;
                border-top: 2px solid #3498db;
                margin: 20px 0;
            }
            '''
        )
        
        # 创建牌组
        deck = genanki.Deck(
            random.randrange(1 << 30, 1 << 31),  # 随机牌组ID
            deck_name
        )
        
        # 添加卡片到牌组
        for card_data in cards_data:
            note = genanki.Note(
                model=model,
                fields=[card_data['front'], card_data['back']]
            )
            deck.add_note(note)
        
        # 生成APKG文件
        package = genanki.Package(deck)
        package.write_to_file(apkg_file_path)
        
        print(f"转换完成！")
        print(f"生成了 {len(cards_data)} 张卡片")
        print(f"APKG文件已保存到: {apkg_file_path}")
        print(f"牌组名称: {deck_name}")
        
        # 显示前几张卡片的预览
        print("\n前3张卡片预览:")
        for i, card in enumerate(cards_data[:3], 1):
            print(f"\n卡片 {i}:")
            print(f"正面: {card['front']}")
            print(f"背面: {card['back'][:100]}{'...' if len(card['back']) > 100 else ''}")
        
        return True
        
    except FileNotFoundError:
        print(f"错误: 找不到文件 {word_file_path}")
        return False
    except Exception as e:
        print(f"处理文件时出错: {str(e)}")
        return False

def advanced_word_to_apkg(word_file_path, apkg_file_path, deck_name="学习卡片"):
    """
    高级版本：支持更复杂的中文数字格式
    支持: 一、二、三、或 （一）（二）（三）等格式
    """
    
    # 更复杂的正则表达式，匹配多种中文数字格式
    patterns = [
        r'^[一二三四五六七八九十][、。]',  # 一、二、三、
        r'^（[一二三四五六七八九十]）',     # （一）（二）（三）
        r'^[一二三四五六七八九十][\s]*$',   # 单独的中文数字
        r'^第[一二三四五六七八九十]+[章节课]',  # 第一章、第二节等
    ]
    
    try:
        print(f"正在读取Word文档 (高级模式): {word_file_path}")
        doc = Document(word_file_path)
        
        cards_data = []
        current_title = ""
        current_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # 检查是否匹配任何标题格式
            is_title = any(re.match(pattern, text) for pattern in patterns)
            
            if is_title:
                # 保存之前的卡片
                if current_title and current_content:
                    content_text = '<br>'.join(current_content).strip()
                    if content_text:
                        cards_data.append({
                            'front': current_title,
                            'back': content_text
                        })
                
                # 开始新标题
                current_title = text
                current_content = []
            else:
                if current_title:
                    current_content.append(text)
        
        # 处理最后一张卡片
        if current_title and current_content:
            content_text = '<br>'.join(current_content).strip()
            if content_text:
                cards_data.append({
                    'front': current_title,
                    'back': content_text
                })
        
        if not cards_data:
            print("错误: 没有找到符合格式的卡片内容")
            return False
        
        # 使用与基础版本相同的模型和牌组创建逻辑
        return create_anki_package(cards_data, apkg_file_path, deck_name)
        
    except Exception as e:
        print(f"高级转换出错: {str(e)}")
        return False

def create_anki_package(cards_data, apkg_file_path, deck_name):
    """
    创建Anki包的辅助函数
    """
    print(f"找到 {len(cards_data)} 张卡片，开始生成APKG文件...")
    
    # 创建模型
    model = genanki.Model(
        random.randrange(1 << 30, 1 << 31),
        '基础卡片模型',
        fields=[
            {'name': '正面'},
            {'name': '背面'},
        ],
        templates=[
            {
                'name': '卡片1',
                'qfmt': '''
                <div class="card">
                    <div class="question">{{正面}}</div>
                </div>
                ''',
                'afmt': '''
                <div class="card">
                    <div class="question">{{正面}}</div>
                    <hr>
                    <div class="answer">{{背面}}</div>
                </div>
                ''',
            },
        ],
        css='''
        .card {
            font-family: "Microsoft YaHei", "Helvetica Neue", Helvetica, Arial, sans-serif;
            font-size: 18px;
            text-align: left;
            color: #333;
            background-color: #fff;
            padding: 20px;
            line-height: 1.6;
        }
        
        .question {
            font-weight: bold;
            font-size: 20px;
            color: #2c3e50;
            margin-bottom: 15px;
        }
        
        .answer {
            font-size: 16px;
            color: #34495e;
            margin-top: 15px;
        }
        
        hr {
            border: none;
            border-top: 2px solid #3498db;
            margin: 20px 0;
        }
        '''
    )
    
    # 创建牌组
    deck = genanki.Deck(
        random.randrange(1 << 30, 1 << 31),
        deck_name
    )
    
    # 添加卡片
    for card_data in cards_data:
        note = genanki.Note(
            model=model,
            fields=[card_data['front'], card_data['back']]
        )
        deck.add_note(note)
    
    # 生成包
    package = genanki.Package(deck)
    package.write_to_file(apkg_file_path)
    
    print(f"转换完成！生成了 {len(cards_data)} 张卡片")
    print(f"APKG文件已保存到: {apkg_file_path}")
    return True

def main():
    """
    主函数 - 支持命令行参数
    """
    
    if len(sys.argv) < 3:
        print("Word文件转Anki APKG格式转换器")
        print("=" * 40)
        print("使用方法:")
        print("python word_to_apkg.py <Word文件路径> <输出APKG文件路径> [牌组名称] [模式]")
        print("\n参数说明:")
        print("  Word文件路径    : 输入的.docx文件")
        print("  输出APKG文件路径 : 输出的.apkg文件")
        print("  牌组名称        : 可选，默认为'学习卡片'")
        print("  模式           : 可选，'advanced'使用高级模式")
        print("\n示例:")
        print("python word_to_apkg.py notes.docx flashcards.apkg")
        print("python word_to_apkg.py notes.docx flashcards.apkg \"我的学习卡片\"")
        print("python word_to_apkg.py notes.docx flashcards.apkg \"高级卡片\" advanced")
        print("\n注意:")
        print("- 需要安装依赖: pip install python-docx genanki")
        print("- Word文档中以中文数字(一、二、三...)开头的行会被识别为标题")
        print("- 标题下方的内容会成为该卡片的答案")
        sys.exit(1)
    
    word_file = sys.argv[1]
    apkg_file = sys.argv[2]
    deck_name = sys.argv[3] if len(sys.argv) > 3 and sys.argv[3] != 'advanced' else "学习卡片"
    use_advanced = len(sys.argv) > 3 and ('advanced' in sys.argv[3:])
    
    # 检查输入文件是否存在
    if not os.path.exists(word_file):
        print(f"错误: 文件不存在 - {word_file}")
        sys.exit(1)
    
    # 确保输出文件有正确的扩展名
    if not apkg_file.endswith('.apkg'):
        apkg_file += '.apkg'
    
    print(f"输入文件: {word_file}")
    print(f"输出文件: {apkg_file}")
    print(f"牌组名称: {deck_name}")
    print(f"模式: {'高级模式' if use_advanced else '标准模式'}")
    print("-" * 40)
    
    # 执行转换
    if use_advanced:
        success = advanced_word_to_apkg(word_file, apkg_file, deck_name)
    else:
        success = word_to_apkg(word_file, apkg_file, deck_name)
    
    if success:
        print(f"\n✅ 转换成功完成！")
        print(f"现在可以将 '{apkg_file}' 导入到Anki中使用了")
    else:
        print(f"\n❌ 转换失败")
        sys.exit(1)

if __name__ == "__main__":
    # 检查依赖
    try:
        import genanki
        from docx import Document
    except ImportError as e:
        print("错误: 缺少必要的依赖包")
        print("请运行以下命令安装:")
        print("pip install python-docx genanki")
        print(f"\n具体错误: {e}")
        sys.exit(1)
    
    main()