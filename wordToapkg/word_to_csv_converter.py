import csv
import re
from docx import Document

def word_to_anki_csv(word_file_path, csv_file_path):
    """
    将Word文件转换为Anki可用的CSV文件
    
    参数:
    word_file_path: Word文件路径
    csv_file_path: 输出CSV文件路径
    """
    
    # 中文数字映射
    chinese_numbers = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    
    # 创建正则表达式模式，匹配以中文数字开头的行
    pattern = r'^[一二三四五六七八九十]'
    
    try:
        # 读取Word文档
        doc = Document(word_file_path)
        
        # 存储卡片数据
        cards = []
        current_title = ""
        current_content = []
        
        # 遍历文档中的每一段
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # 跳过空行
            if not text:
                continue
            
            # 检查是否是标题行（以中文数字开头）
            if re.match(pattern, text):
                # 如果之前有标题和内容，保存为一张卡片
                if current_title and current_content:
                    content_text = '\n'.join(current_content).strip()
                    if content_text:  # 确保内容不为空
                        cards.append([current_title, content_text])
                
                # 开始新的标题
                current_title = text
                current_content = []
            else:
                # 这是内容行，添加到当前内容中
                if current_title:  # 只有在有标题的情况下才添加内容
                    current_content.append(text)
        
        # 处理最后一张卡片
        if current_title and current_content:
            content_text = '\n'.join(current_content).strip()
            if content_text:
                cards.append([current_title, content_text])
        
        # 写入CSV文件
        with open(csv_file_path, 'w', newline='', encoding='utf-8-sig') as csvfile:
            writer = csv.writer(csvfile)
            
            # 写入表头（可选）
            # writer.writerow(['问题', '答案'])
            
            # 写入卡片数据
            for card in cards:
                writer.writerow(card)
        
        print(f"转换完成！共生成 {len(cards)} 张卡片")
        print(f"CSV文件已保存到: {csv_file_path}")
        
        # 显示前几张卡片的预览
        if cards:
            print("\n前3张卡片预览:")
            for i, card in enumerate(cards[:3], 1):
                print(f"\n卡片 {i}:")
                print(f"标题: {card[0]}")
                print(f"内容: {card[1][:100]}{'...' if len(card[1]) > 100 else ''}")
    
    except FileNotFoundError:
        print(f"错误: 找不到文件 {word_file_path}")
    except Exception as e:
        print(f"处理文件时出错: {str(e)}")

def main():
    """
    主函数 - 支持命令行参数
    """
    import sys
    
    # 检查命令行参数
    if len(sys.argv) != 3:
        print("使用方法:")
        print("python word_to_anki.py <输入的Word文件路径> <输出的CSV文件路径>")
        print("\n示例:")
        print("python word_to_anki.py my_notes.docx anki_cards.csv")
        print("python word_to_anki.py /path/to/document.docx /path/to/output.csv")
        sys.exit(1)
    
    word_file = sys.argv[1]
    csv_file = sys.argv[2]
    
    print(f"输入文件: {word_file}")
    print(f"输出文件: {csv_file}")
    print("开始转换...")
    
    # 执行转换
    word_to_anki_csv(word_file, csv_file)

if __name__ == "__main__":
    # 安装所需依赖的命令:
    # pip install python-docx
    
    main()

# 高级版本：支持更复杂的中文数字格式
def advanced_word_to_anki_csv(word_file_path, csv_file_path):
    """
    支持更复杂中文数字格式的转换函数
    支持: 一、二、三、或 （一）（二）（三）等格式
    """
    
    # 更复杂的正则表达式，匹配多种中文数字格式
    patterns = [
        r'^[一二三四五六七八九十][、。]',  # 一、二、三、
        r'^（[一二三四五六七八九十]）',     # （一）（二）（三）
        r'^[一二三四五六七八九十][\s]*$',   # 单独的中文数字
    ]
    
    try:
        doc = Document(word_file_path)
        cards = []
        current_title = ""
        current_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # 检查是否匹配任何标题格式
            is_title = any(re.match(pattern, text) for pattern in patterns)
            
            if is_title:
                # 保存之前的卡片
                if current_title and current_content:
                    content_text = '\n'.join(current_content).strip()
                    if content_text:
                        cards.append([current_title, content_text])
                
                # 开始新标题
                current_title = text
                current_content = []
            else:
                if current_title:
                    current_content.append(text)
        
        # 处理最后一张卡片
        if current_title and current_content:
            content_text = '\n'.join(current_content).strip()
            if content_text:
                cards.append([current_title, content_text])
        
        # 写入CSV
        with open(csv_file_path, 'w', newline='', encoding='utf-8-sig') as csvfile:
            writer = csv.writer(csvfile)
            for card in cards:
                writer.writerow(card)
        
        print(f"高级转换完成！共生成 {len(cards)} 张卡片")
        
    except Exception as e:
        print(f"高级转换出错: {str(e)}")

# 使用高级版本的示例
# advanced_word_to_anki_csv("input.docx", "anki_cards_advanced.csv")