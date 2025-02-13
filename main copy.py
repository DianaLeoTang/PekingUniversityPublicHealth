# from docx import Document
# import re

# def process_word_document(input_path, output_path):
#     """
#     处理Word文档中的不必要折行
    
#     Args:
#         input_path: 输入Word文档的路径
#         output_path: 输出Word文档的路径
#     """
#     # 加载Word文档
#     doc = Document(input_path)
    
#     # 中文标点符号集合
#     punctuation = {'。', '，', '！', '？', '；', '：', '、', '"', '"', ''', ''', '【', '】', '（', '）', '《', '》', '—', '…', '．'}
    
#     # 处理每个段落
#     for paragraph in doc.paragraphs:
#         if not paragraph.text.strip():  # 跳过空段落
#             continue
            
#         # 获取段落文本
#         text = paragraph.text
        
#         # 分割成行
#         lines = text.split('\n')
        
#         # 处理每一行
#         result_lines = []
#         temp_line = ''
        
#         for line in lines:
#             line = line.strip()  # 移除首尾空白
#             if not line:  # 跳过空行
#                 if temp_line:
#                     result_lines.append(temp_line)
#                     temp_line = ''
#                 continue
                
#             # 如果当前行末尾有标点符号
#             if line and line[-1] in punctuation:
#                 if temp_line:
#                     temp_line += line
#                     result_lines.append(temp_line)
#                     temp_line = ''
#                 else:
#                     result_lines.append(line)
#             else:
#                 # 如果没有标点符号，累积到临时行
#                 temp_line += line
        
#         # 处理最后一行
#         if temp_line:
#             result_lines.append(temp_line)
        
#         # 更新段落文本
#         new_text = '\n'.join(result_lines)
#         paragraph.text = new_text
        
#         # 保持原始格式
#         for run in paragraph.runs:
#             # 这里可以添加格式保持的代码，如字体、大小等
#             pass
    
#     # 保存修改后的文档
#     doc.save(output_path)

# # 使用示例
# if __name__ == "__main__":
#     input_file = "./我的科研助理：GPT.docx"  # 输入文件路径
#     output_file = "./output.docx"  # 输出文件路径
#     process_word_document(input_file, output_file)

# from docx import Document
# import re

# def remove_unnecessary_line_breaks(doc):
#     # 初始化一个空列表来存储处理后的段落文本
#     processed_paragraphs = []

#     # 遍历文档中的每一个段落
#     for paragraph in doc.paragraphs:
#         text = paragraph.text
#         # 将段落文本按行分割
#         lines = text.split('\n')
#         # 初始化一个空列表来存储处理后的行
#         processed_lines = []
#         # 遍历每一行
#         for i in range(len(lines)):
#             # 如果当前行不是最后一行，并且当前行末尾没有标点符号
#             if i < len(lines) - 1 and not re.search(r'[。！？，、；：,.!?;:]$', lines[i]):
#                 # 将当前行与下一行合并，并去除多余的空白字符
#                 lines[i + 1] = lines[i].rstrip() + ' ' + lines[i + 1].lstrip()
#             else:
#                 # 否则，直接将当前行添加到处理后的行列表中
#                 processed_lines.append(lines[i])
#         # 将处理后的行重新组合成段落文本
#         processed_text = ' '.join(processed_lines)
#         processed_paragraphs.append(processed_text)

#     # 创建一个新的Word文档
#     new_doc = Document()
#     # 将处理后的段落添加到新文档中
#     for paragraph in processed_paragraphs:
#         new_doc.add_paragraph(paragraph)

#     return new_doc

# # 读取原始Word文档
# input_doc_path = './我的科研助理：GPT.docx'  # 替换为你的输入文件路径
# doc = Document(input_doc_path)

# # 处理文档
# new_doc = remove_unnecessary_line_breaks(doc)

# # 保存处理后的文档
# output_doc_path = 'output.docx'  # 替换为你的输出文件路径
# new_doc.save(output_doc_path)

# print(f"处理后的文档已保存到: {output_doc_path}")

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
import re

def is_line_break(element):
    """检查一个XML元素是否为软回车（<w:br/>）"""
    # 确保 element 是 OxmlElement 类型，并且其 tag 属性以 'br' 结尾
    return isinstance(element, OxmlElement) and element.tag.endswith('br')

def remove_unnecessary_line_breaks(doc):
    """删除段落中不必要的软回车（合并无标点结尾的行）"""
    for para in doc.paragraphs:
        for run in para.runs:
            # 获取 run 中的所有 XML 元素
            for elem in run._r:
                # 检查是否是折行符
                if is_line_break(elem):
                    run._r.remove(elem)  # 移除折行符
    return doc

# 读取输入文档
input_doc = Document("./我的科研助理：GPT.docx")

# 处理文档
processed_doc = remove_unnecessary_line_breaks(input_doc)

# 保存输出
processed_doc.save("./output.docx")